"""Generate the outreach flow document (flow scenario docx)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

GREY = RGBColor(0x55, 0x55, 0x55)


def add_para(doc, text, bold=False, size=10.5, color=None, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(text, style='List Bullet')
    for r in p.runs:
        r.font.size = Pt(size)
    return p


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
    return t


doc = Document()
doc.core_properties.title = 'Multi-Number Outreach — Flow & Throughput Guide'
doc.core_properties.author = 'OpenWA'

title = doc.add_heading('Multi-Number Cold Outreach — Message Flow & Throughput', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = add_para(doc, 'How a wave actually flows from the API through N sessions, with the randomized '
                'cooldown (4–8 min) technique. Includes the 10-aged-numbers = 10,000 msgs math.',
               size=10, color=GREY, space_after=10)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

# =====================================================================
add_para(doc, '1. How one wave flows, end to end', bold=True, size=13)
steps = [
    '1. Create the wave — send the contact list, the message template, the session pool (by name), and the strategy.',
    '2. Round-robin allocation (pure math, no sending yet) — contacts are balanced across sessions so each gets an equal share (within ±1).',
    '3. Warm-up cap — each session is capped at its account-age allowance (20 → 40 → 80 → 160 → 320 → 640 → 1000/day). Mature numbers = 1000/day.',
    '4. Burst slicing — each session’s share is split into bursts of `burstSize`, with a RANDOM cooldown (default 4–8 min) drawn after each burst.',
    '5. Dispatch loop — every 2s the orchestrator checks each session: READY? not restricted? cooled down? has a burst left? If yes → submit that burst.',
    '6. Per-send humanized flow (inside each burst) — save contact → pre-check number on WhatsApp → typing indicator → send → humanized 30–120s delay → next.',
    '7. Progress — per-session sent/failed/pending tally is persisted; a wave completes when every session finishes its bursts.',
]
for s in steps:
    add_bullet(doc, s)

# =====================================================================
add_para(doc, '2. The per-send technique (the ~35s flow)', bold=True, size=13)
add_table(doc,
    ['Step', 'Action', 'Approx time'],
    [
        ['1', 'saveContactFirst — upsert into address book', '1–2s'],
        ['2', 'preCheckNumbers — verify registered on WhatsApp', '1–2s'],
        ['3', 'typing indicator', '1–2s'],
        ['4', 'send message', '1–3s'],
        ['5', 'humanized random delay between sends', '25–35s (avg)'],
        ['', 'Total per contact', '≈ 30–35s avg'],
    ])
add_para(doc, 'Note: sends are serial per number (one phone = one engine). A single number cannot fire '
         'two messages at the same instant — that is bot behaviour and is exactly what gets flagged.',
         size=9.5, color=GREY, space_after=8)

# =====================================================================
add_para(doc, '3. Randomized cooldown between bursts (4–8 min)', bold=True, size=13)
add_para(doc, 'After each burst the session pauses a RANDOM 4–8 minutes before the next burst. The random '
         'value makes the rhythm look human instead of metronomic:', size=10.5)
add_bullet(doc, 'burst 1: send 20 → wait ~4.2 min')
add_bullet(doc, 'burst 2: send 20 → wait ~6.7 min')
add_bullet(doc, 'burst 3: send 20 → wait ~5.0 min')
add_bullet(doc, 'burst 4: send 20 → wait ~7.5 min … and so on')
add_table(doc,
    ['Strategy knob', 'Purpose', 'Value for 10k/day'],
    [
        ['burstSize', 'Recipients per burst before cooldown', '20–30 (not 10)'],
        ['cooldownMinMs', 'Cooldown lower bound', '240000 ms (4 min)'],
        ['cooldownMaxMs', 'Cooldown upper bound', '480000 ms (8 min)'],
        ['minDelayMs / maxDelayMs', 'Humanized delay between sends', '20000 / 50000 (avg ~35s)'],
        ['saveContactFirst / preCheckNumbers', 'Anti-ban addressbook + WhatsApp check', 'true'],
        ['warmupSchedule', 'Max sends/day by account age', '[20,40,80,160,320,640,1000]'],
    ])

# =====================================================================
add_para(doc, '4. Scenario A — ONE aged number (2–3 months)', bold=True, size=13)
add_para(doc, 'Single mature session, window 6am–10pm IST (16h), 30–120s default pacing vs the ~35s technique:',
         size=10.5)
add_table(doc,
    ['Metric', '30–120s pacing (default)', '~35s technique (20-msg bursts, 4–8min cool)'],
    [
        ['Budget per message', '75s avg', '35s avg'],
        ['Capacity in 16h', '~760 msgs', '1,000 msgs ✓'],
        ['Practical new-chat daily (safe)', '100–150', '100–150'],
        ['Time to clear 1,000 contacts', '—', '~14.6h'],
    ])
add_para(doc, 'Conclusion: ONE number cannot safely clear a large list in a single day. 1,000/day is the '
         'absolute pacing ceiling; the SAFE cold-outreach figure stays ~100–150 new chats/day per number.',
         size=9.5, color=GREY, space_after=8)

# =====================================================================
add_para(doc, '5. Scenario B — FIVE aged numbers', bold=True, size=13)
add_table(doc,
    ['Setup per number', 'Per number', '5 numbers total', 'Fits 16h?'],
    [
        ['Safe pace ~150/day', '150 msgs / ~1.6h', '750 msgs', 'Yes, huge slack'],
        ['~35s + 20/burst + 4–8min cool', '1,000 msgs / ~14.6h', '5,000 msgs', 'Yes'],
        ['Max pacing, no cooldown', '1,000 msgs / ~9.7h', '5,000 msgs', 'Yes (~6h slack)'],
    ])
add_para(doc, 'Five mature numbers at the full technique clear 5,000 contacts in one 16h day.',
         size=9.5, color=GREY, space_after=8)

# =====================================================================
add_para(doc, '6. Scenario C — TEN aged numbers → 10,000 msgs', bold=True, size=13)
add_para(doc, 'The math that makes 10,000 in 16h possible:', size=10.5)
add_bullet(doc, '10,000 ÷ 10 numbers = 1,000 messages per number.')
add_bullet(doc, '16h = 57,600s → budget = 57.6s per message per number.')
add_bullet(doc, '~35s avg flow × 1,000 = ~9.7h of pure sending — inside the 16h window with ~6.3h of cooldown budget.')
add_bullet(doc, 'Cooldown total with 20-msg bursts (50 bursts): 49 × ~6min ≈ 4.9h.')
add_bullet(doc, 'Total per number ≈ 14.6h → comfortably inside 16h across the whole 16h window.')
add_table(doc,
    ['burstSize', 'Bursts/number', 'Send time (1000×35s)', 'Cool time (avg 6min)', 'Total/number', 'Fits 16h?'],
    [
        ['10', '100', '9.7h', '99×6min = 9.9h', '~19.6h', 'NO'],
        ['20', '50', '9.7h', '49×6min = 4.9h', '~14.6h', 'YES ✓'],
        ['25', '40', '9.7h', '39×6min = 3.9h', '~13.6h', 'YES ✓'],
        ['30', '34', '9.7h', '33×6min = 3.3h', '~13.0h', 'YES ✓'],
    ])
add_para(doc, 'Fleet throughput (burstSize 20): ~685 msgs/hr across 10 numbers (~68.5/number/hr). '
         '10,000 contacts dispatched in ~14.6h of a 16h IST window, with all sessions running in parallel '
         'and separate random 4–8 min cool phases.', size=10.5, space_after=8)

# =====================================================================
add_para(doc, '7. The one honest warning', bold=True, size=13)
add_para(doc, '1,000 pure cold messages per number per day is ~7–10× above the community-safe zone '
         '(~100–150 new chats/day), regardless of account age. The math works — WhatsApp tolerance is the '
         'unknown. If you run a 10k/day wave: spread the waves with bursts + cooldowns exactly as above, '
         'funnel toward people who reply, and run with an automatic stop the instant any session reports a '
         'restriction.', size=10.5, space_after=8)

doc.save('OpenWA-MultiLine-Outreach-Flow-10k.docx')
print('saved OpenWA-MultiLine-Outreach-Flow-10k.docx')